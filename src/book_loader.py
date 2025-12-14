import fitz  # PyMuPDF
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
import docx
import re
from dataclasses import dataclass
from typing import List, Set, Tuple


@dataclass
class Chapter:
    title: str
    content: str


class BookLoader:
    # 垃圾信息关键词（版权页、出版社信息等）
    METADATA_KEYWORDS = [
        "ISBN", "Copyright", "版权", "出版社", "出版年", "印刷", "定价",
        "CIP", "All rights reserved", "Printed in", "责任编辑", "装帧设计",
        "字数", "印张", "版次", "书号", "经销", "开本"
    ]

    # 增强的章节标题匹配模式
    CHAPTER_PATTERNS = [
        r'(^第[0-9一二三四五六七八九十百千]+[章|节|回|卷|部|篇].*)',  # 标准：第1章 / 第一卷
        r'(^Chapter\s+\d+.*)',  # 英文：Chapter 1
        r'(^Part\s+[One|Two|Three|I|II|III|IV|V|VI].*)',  # 英文：Part One
        r'(^[0-9一二三四五六七八九十百]+\s+[\u4e00-\u9fa5]{2,})',  # 纯数字+文字：1. 开始
        r'(^前言|^序言|^引子|^楔子|^尾声|^后记|^番外|^终章|^结语|^Prologue|^Epilogue|^Introduction|^Preface)',  # 特殊结构
        r'(^目录|^Table of Contents)',  # 目录页本身
        r'(^作者.*|^致谢.*)'  # 匹配 "作者的话", "致谢" 等
    ]

    @staticmethod
    def load_book(file) -> List[Chapter]:
        """工厂方法：根据文件后缀分发处理逻辑"""
        filename = file.name.lower()
        chapters = []

        try:
            if filename.endswith('.epub'):
                chapters = BookLoader._parse_epub(file)
            elif filename.endswith('.docx'):
                chapters = BookLoader._parse_docx(file)
            elif filename.endswith('.pdf'):
                chapters = BookLoader._parse_pdf(file)
            elif filename.endswith('.txt'):
                chapters = BookLoader._parse_txt(file)
            else:
                raise ValueError("不支持的文件格式")

            # 统一进行垃圾章节过滤
            return BookLoader._filter_junk_chapters(chapters)
        except Exception as e:
            # 捕获解析错误，避免整个程序崩溃
            print(f"解析书籍出错: {e}")
            raise e

    @staticmethod
    def _is_chapter_title(text: str, custom_titles: Set[str] = None) -> bool:
        """判断一行文本是否像章节标题"""
        text = text.strip()
        if not text:
            return False

        # 1. 优先匹配动态提取的目录标题 (精确匹配)
        if custom_titles and text in custom_titles:
            return True

        # 2. 长度过滤 (标题一般不会太长)
        if len(text) > 50:
            return False

        # 3. 正则匹配通用模式
        for pattern in BookLoader.CHAPTER_PATTERNS:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        return False

    @staticmethod
    def _extract_toc_titles(content: str) -> Set[str]:
        """目录嗅探：尝试从文本开头识别文本目录，提取短标题"""
        found_titles = set()
        lines = content.split('\n')

        scan_limit = min(len(lines), 300)  # 稍微增加扫描行数
        in_toc_area = False

        for i in range(scan_limit):
            line = lines[i].strip()
            if not line:
                continue

            if re.match(r'^目录$|^Table of Contents$', line, re.IGNORECASE):
                in_toc_area = True
                continue

            if in_toc_area:
                if len(line) > 25:  # 遇到长句认为目录结束
                    break
                if not line.isdigit() and "..." not in line:  # 过滤页码和省略号
                    found_titles.add(line)
                if len(found_titles) > 30:
                    break

        return found_titles

    @staticmethod
    def _filter_junk_chapters(chapters: List[Chapter]) -> List[Chapter]:
        """过滤掉版权页、目录页等非正文内容"""
        clean_chapters = []
        for chap in chapters:
            title_lower = chap.title.lower()
            content_lower = chap.content.lower()[:500]

            # 1. 检查标题是否包含明显的垃圾关键词
            if any(kw in title_lower for kw in ["版权", "copyright", "colophon", "table of contents"]):
                continue
            if title_lower.strip() == "目录":
                continue

            # 2. 检查内容密度 (放宽限制，避免误杀短文)
            if len(chap.content) < 500:  # 稍微降低阈值
                hit_count = sum(1 for kw in BookLoader.METADATA_KEYWORDS if kw.lower() in content_lower)
                if hit_count >= 2 or ("isbn" in content_lower):
                    continue

            clean_chapters.append(chap)

        return clean_chapters if clean_chapters else chapters

    # --- EPUB 解析逻辑重构 (核心修改) ---
    @staticmethod
    def _parse_epub(file) -> List[Chapter]:
        temp_path = f"temp_{file.name}"
        with open(temp_path, "wb") as f:
            f.write(file.read())

        book = epub.read_epub(temp_path)
        chapters = []

        # 1. 尝试从 NCX/Nav 目录读取 (最准确)
        toc_items = BookLoader._flatten_epub_toc(book.toc)

        if toc_items:
            # 如果有目录，按照目录抓取
            for link in toc_items:
                try:
                    # link.href 可能是 'chap1.xhtml' 或 'chap1.xhtml#anchor'
                    href_parts = link.href.split('#')
                    file_href = href_parts[0]
                    # anchor = href_parts[1] if len(href_parts) > 1 else None

                    # 根据 href 找到 item
                    item = book.get_item_with_href(file_href)
                    if item:
                        soup = BeautifulSoup(item.get_content(), 'html.parser')
                        # 简单处理：提取该文件的全部文字
                        # (精细化处理 anchor 比较复杂，这里先确保整章能读到)
                        text = soup.get_text().strip()
                        if len(text) > 50:
                            chapters.append(Chapter(title=link.title, content=text))
                except Exception:
                    continue
        else:
            # 2. 目录为空，回退到 Spine (阅读顺序) 遍历
            # Spine 是书籍定义的线性阅读顺序，比 get_items() 靠谱
            for item_id, linear in book.spine:
                item = book.get_item_with_id(item_id)
                if item:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    text = soup.get_text().strip()

                    # 尝试寻找标题
                    title = ""
                    header = soup.find(['h1', 'h2', 'h3'])
                    if header:
                        title = header.get_text().strip()
                    else:
                        # 如果没有标题，尝试用文件名或 ID
                        title = item.get_name()

                    if len(text) > 50:
                        chapters.append(Chapter(title=title, content=text))

        return chapters

    @staticmethod
    def _flatten_epub_toc(toc, depth=0):
        """递归展平 EPUB 的嵌套目录"""
        items = []
        for item in toc:
            if isinstance(item, epub.Link):
                items.append(item)
            elif isinstance(item, (tuple, list)):
                # 目录项可能是 (Section, [Children])
                section, children = item
                if isinstance(section, epub.Link):
                    items.append(section)
                items.extend(BookLoader._flatten_epub_toc(children, depth + 1))
        return items

    @staticmethod
    def _parse_docx(file) -> List[Chapter]:
        doc = docx.Document(file)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        toc_titles = BookLoader._extract_toc_titles(full_text)

        chapters = []
        current_title = "正文"
        current_content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            is_heading = para.style.name.startswith('Heading')
            is_pattern_match = BookLoader._is_chapter_title(text, toc_titles)

            if is_heading or is_pattern_match:
                if current_content:
                    chapters.append(Chapter(title=current_title, content="\n".join(current_content)))
                current_title = text
                current_content = []
            else:
                current_content.append(text)

        if current_content:
            chapters.append(Chapter(title=current_title, content="\n".join(current_content)))

        return chapters

    @staticmethod
    def _parse_pdf(file) -> List[Chapter]:
        doc = fitz.open(stream=file.read(), filetype="pdf")
        toc = doc.get_toc()
        chapters = []

        if toc:
            for i in range(len(toc)):
                title = toc[i][1]
                start_page = toc[i][2] - 1
                end_page = toc[i + 1][2] - 1 if i + 1 < len(toc) else doc.page_count

                text = ""
                for page_num in range(start_page, end_page):
                    text += doc[page_num].get_text()

                if text.strip():
                    chapters.append(Chapter(title=title, content=text))
        else:
            full_text = ""
            for page in doc:
                full_text += page.get_text()
            chapters = BookLoader._split_text_by_patterns(full_text)

        return chapters

    @staticmethod
    def _parse_txt(file) -> List[Chapter]:
        content = file.read().decode('utf-8')
        return BookLoader._split_text_by_patterns(content)

    @staticmethod
    def _split_text_by_patterns(content: str) -> List[Chapter]:
        toc_titles = BookLoader._extract_toc_titles(content)
        lines = content.split('\n')
        chapters = []
        current_title = "正文"
        current_content = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            if BookLoader._is_chapter_title(line, toc_titles):
                if current_content:
                    chapters.append(Chapter(title=current_title, content="\n".join(current_content)))
                current_title = line
                current_content = []
            else:
                current_content.append(line)

        if current_content:
            chapters.append(Chapter(title=current_title, content="\n".join(current_content)))

        return chapters
